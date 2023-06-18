import datetime
import json
import subprocess

try:
    from docx import Document
except ModuleNotFoundError:
    print("docx not installed\ninstalling python-docx...")
    subprocess.check_call(["pip", "install", "python-docx"])
    print("successfully installed python-docx\nimporting python-docx now")
    from docx import Document


class TrainingReportEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, datetime.datetime):
            return obj.isoformat()
        return super().default(obj)


class TrainingReportDecoder(json.JSONDecoder):
    def __init__(self, *args, **kwargs):
        super().__init__(object_hook=self.object_hook, *args, **kwargs)

    @staticmethod
    def object_hook(dct):
        if "__class__" in dct:
            class_name = dct.pop("__class__")
            if class_name == "training_report":
                return training_report()
        return dct


class tuple1:
    __max_length = 50
    __text: str
    __hours: int

    def __init__(self, text: str, hours: int):
        self.__text = text[:self.__max_length]
        self.__hours = hours

    def get_text(self) -> str:
        return self.__text

    def get_hours(self) -> int:
        return self.__hours

    @staticmethod
    def generate_str_from_list_text(array: list) -> str:
        gen_text = ""
        for obj in array:
            gen_text = gen_text + obj.get_text() + "/n"
        return gen_text

    @staticmethod
    def generate_str_from_list_hours(array: list) -> str:
        gen_text = ""
        for obj in array:
            gen_text = gen_text + obj.get_hours() + "/n"
        return gen_text


class training_report:
    # $YOT      Year of training
    __yot: int

    # $WOTB     Week of training Beginning
    __wotb: datetime

    # $WOTE     Week of training Ending
    __wote: datetime

    # $OA       Operational activities
    __oa: list[tuple1]

    # $I        Instructions
    __i: list[tuple1]

    # $TST      Topics of school teaching
    __tst: list[tuple1]

    __doc: Document()
    __training_begin: datetime

    def __init__(self):
        self.__doc = Document("BerichtVorlage.docx")
        self.__training_begin = datetime.datetime(2022, 8, 1)
        self.set_head_table(f"{datetime.datetime.now().strftime('%Y')}-W{datetime.datetime.now().strftime('%W')}")
        self.__oa = []
        self.__i = []
        self.__tst = []

    def to_json(self) -> str:
        json_dict = {
            "yot": self.__yot,
            "wotb": self.__wotb.isoformat(),
            "wote": self.__wote.isoformat(),
            "oa": [obj.__dict__ for obj in self.__oa],
            "i": [obj.__dict__ for obj in self.__i],
            "tst": [obj.__dict__ for obj in self.__tst]
        }
        return json.dumps(json_dict)

    @staticmethod
    def from_json(json_str: str) -> 'training_report':
        json_data = json.loads(json_str)
        report = training_report()
        report.__yot = json_data["yot"]
        report.__wotb = datetime.datetime.strptime(json_data["wotb"], "%Y-%m-%dT%H:%M:%S")
        report.__wote = datetime.datetime.strptime(json_data["wote"], "%Y-%m-%dT%H:%M:%S")

        report.__oa = []
        for item in json_data["oa"]:
            obj = tuple1(item["_tuple1__text"], item["_tuple1__hours"])
            report.__oa.append(obj)

        report.__i = []
        for item in json_data["i"]:
            obj = tuple1(item["_tuple1__text"], item["_tuple1__hours"])
            report.__i.append(obj)

        report.__tst = []
        for item in json_data["tst"]:
            obj = tuple1(item["_tuple1__text"], item["_tuple1__hours"])
            report.__tst.append(obj)

        return report

    def __calc_abj(self, wotb: datetime) -> int:
        diff = wotb.year - self.__training_begin.year
        if wotb.isocalendar()[1] > self.__training_begin.isocalendar()[1]:
            diff -= 1
        return diff

    def set_head_table(self, cw: str):
        # kw must be format like jjjj-Www Example: "2023-W26"
        self.__wotb = datetime.datetime.strptime(cw + '-1', "%Y-W%W-%w")
        self.__wote = datetime.datetime.strptime(cw + '-5', "%Y-W%W-%w")
        self.__yot = self.__calc_abj(self.__wotb)

    def add_oa(self, text: str, hours: int):
        if len(self.__oa) <= 8:
            self.__oa.append(tuple1(text, hours))

    def remove_oa(self, index: int):
        self.__oa.pop(index)

    def add_i(self, text: str, hours: int):
        if len(self.__oa) <= 8:
            self.__i.append(tuple1(text, hours))

    def remove_i(self, index: int):
        self.__i.pop(index)

    def add_tst(self, text: str, hours: int):
        if len(self.__oa) <= 8:
            self.__tst.append(tuple1(text, hours))

    def remove_tst(self, index: int):
        self.__tst.pop(index)

    def __replace_markers(self, text: str) -> str:
        if text.startswith("$YOT"):
            return str(self.__yot)
        elif text.startswith("$WOTB"):
            return self.__wotb.strftime("%d.%m.%Y")
        elif text.startswith("$WOTE"):
            return self.__wote.strftime("%d.%m.%Y")
        elif text.startswith("$OA"):
            return tuple1.generate_str_from_list_text(self.__oa)
        elif text.startswith("$OAH"):
            return tuple1.generate_str_from_list_hours(self.__oa)
        elif text.startswith("$I"):
            return tuple1.generate_str_from_list_text(self.__i)
        elif text.startswith("$IH"):
            return tuple1.generate_str_from_list_hours(self.__i)
        elif text.startswith("$TST"):
            return tuple1.generate_str_from_list_text(self.__tst)
        elif text.startswith("$TSTH"):
            return tuple1.generate_str_from_list_hours(self.__tst)

    def check_work_hours(self) -> (bool, int):
        work_hours = 0

        for obj in self.__oa:
            work_hours = work_hours + obj.get_hours()
        for obj in self.__i:
            work_hours = work_hours + obj.get_hours()
        for obj in self.__tst:
            work_hours = work_hours + obj.get_hours()

        if work_hours == 40:
            return True, work_hours
        else:
            return False, work_hours

    def save_document_to(self, filename: str):
        for table in self.__doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.startswith("$"):
                        cell.text = self.__replace_markers(cell.text)
        self.__doc.save(filename)

    """
    const string TOP_LEFT_JOINT = "┌";
    const string TOP_RIGHT_JOINT = "┐";
    const string BOTTOM_LEFT_JOINT = "└";
    const string BOTTOM_RIGHT_JOINT = "┘";
    const string TOP_JOINT = "┬";
    const string BOTTOM_JOINT = "┴";
    const string LEFT_JOINT = "├";
    const string JOINT = "┼";
    const string RIGHT_JOINT = "┤";
    const char HORIZONTAL_LINE = '─';
    const char PADDING = ' ';
    const string VERTICAL_LINE = "│";
    """

    def print_haed(self):
        print("┌" + "─Week of training Beginning──" +
              "┬" + "─Week of training Ending──" +
              "┬" + "─Year of training──" + "┐")
        print("│" + " " * 10 + self.__wotb.strftime("%d.%m.%Y") + " " * 9 +
              "│" + " " * 8 + self.__wote.strftime("%d.%m.%Y") + " " * 8 +
              "│" + " " * 9 + str(self.__yot) + " " * 9 + "│")
        print("└" + "─" * 29 + "┴" + "─" * 26 + "┴" + "─" * 19 + "┘")

    def print_tables(self):
        print("┌" + "─" + "Operational activities" + "─" * 27 + "┬──┐")
        for tup in self.__oa:
            print("│" + tup.get_text() + (" " * (50 - len(tup.get_text()))) + "│" +
                  (" " * (2 - len(str(tup.get_hours())))) + str(tup.get_hours()) + "│")
        print("└" + "─" * 50 + "┴──┘")
        print("┌" + "─" + "Instructions" + "─" * 37 + "┬──┐")
        for tup in self.__i:
            print("│" + tup.get_text() + (" " * (50 - len(tup.get_text()))) + "│" +
                  (" " * (2 - len(str(tup.get_hours())))) + str(tup.get_hours()) + "│")
        print("└" + "─" * 50 + "┴──┘")
        print("┌" + "─" + "Topics of school teaching" + "─" * 24 + "┬──┐")
        for tup in self.__tst:
            print("│" + tup.get_text() + (" " * (50 - len(tup.get_text()))) + "│" +
                  (" " * (2 - len(str(tup.get_hours())))) + str(tup.get_hours()) + "│")
        print("└" + "─" * 50 + "┴──┘")

    def print_all(self):
        self.print_haed()
        self.print_tables()

    def print_doc_paragraphs(self):
        for p in self.__doc.paragraphs:
            print("-" * 32 + "\nNew Paragraph\n" + "-" * 32)
            print(p.text)

    def print_doc_tables(self):
        for table in self.__doc.tables:
            print("-" * 32 + "\nNew Table\n" + "-" * 32)
            for row in table.rows:
                print("|".join([cell.text for cell in row.cells]))

    def get_paragraphs(self):
        return self.__doc.paragraphs

    def get_tables(self):
        return self.__doc.tables


def exit_app(training_rep: training_report):
    json_str = training_rep.to_json()
    with open("training_report.json", "w") as file:
        file.write(json_str)
    exit()


def main():
    # Beispiel für das Schreiben in eine JSON-Datei
    abb = training_report()
    abb.print_all()
    abb.set_head_table("2023-W13")

    abb.add_oa("RMM", 12)
    abb.add_oa("Außendienst", 8)

    exit_app(abb)

    # Beispiel für das Lesen aus einer JSON-Datei
    with open("training_report.json", "r") as file:
        json_str = file.read()
    abb = training_report.from_json(json_str)
    abb.print_all()


if __name__ == '__main__':
    main()
