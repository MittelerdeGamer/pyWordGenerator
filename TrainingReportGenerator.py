import datetime
import json
import random
import subprocess
import os

#
# last tr kw25
#
# missing kw14
# missing kw15
# missing kw16
#

try:
    from docx import Document
except ModuleNotFoundError:
    print("docx not installed\ninstalling python-docx...")
    subprocess.check_call(["pip", "install", "python-docx"])
    print("successfully installed python-docx\nimporting python-docx now")
    from docx import Document

try:
    from termcolor import colored
except ModuleNotFoundError:
    print("termcolor not installed\ninstalling termcolor...")
    subprocess.check_call(["pip", "install", "termcolor"])
    print("successfully installed termcolor\nimporting termcolor now")
    from termcolor import colored

document_break_line = """
"""


class TrainingReportEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, datetime.datetime):
            return obj.isoformat()
        return super().default(obj)


class TrainingReportDecoder(json.JSONDecoder):
    def __init__(self, *args, **kwargs):
        super().__init__(object_hook=self.object_hook, *args, **kwargs)

    @staticmethod
    def object_hook(dct):
        if "__class__" in dct:
            class_name = dct.pop("__class__")
            if class_name == "training_report":
                return training_report()
        return dct


class tuple1:
    __max_length = 65
    __text: str
    __hours: int

    def __init__(self, text: str, hours: int):
        self.__text = text[:self.__max_length]
        if 1 <= hours <= 40:
            self.__hours = hours
        else:
            if hours < 1:
                self.__hours = 1
            elif 40 < hours:
                self.__hours = 40
            print(colored("ValueWarning: Hours must be between 1 and 40.\n", "yellow"))

    def get_text(self) -> str:
        return self.__text

    def get_hours(self) -> int:
        return self.__hours

    def set_hours(self, hours: int):
        if 1 <= hours <= 40:
            self.__hours = hours
        else:
            if hours < 1:
                self.__hours = 1
            elif 40 < hours:
                self.__hours = 40
            print(colored("ValueWarning: Hours must be between 1 and 40.\n", "yellow"))

    @staticmethod
    def generate_str_from_list_text(array: list) -> str:
        gen_text = ""
        for obj in array:
            gen_text = gen_text + obj.get_text() + document_break_line
        return gen_text

    @staticmethod
    def generate_str_from_list_hours(array: list) -> str:
        gen_text = ""
        for obj in array:
            gen_text = gen_text + str(obj.get_hours()) + document_break_line
        return gen_text


class training_report:
    # $YOT      Year of training
    __yot: int

    # $WOTB     Week of training Beginning
    __wotb: datetime

    # $WOTE     Week of training Ending
    __wote: datetime

    # $OA       Operational activities
    __oa: list[tuple1]

    # $I        Instructions
    __i: list[tuple1]

    # $TST      Topics of school teaching
    __tst: list[tuple1]

    __doc: Document()
    __training_begin: datetime

    def __init__(self):
        self.__doc = Document("BerichtVorlage.docx")
        self.__training_begin = datetime.datetime(2022, 8, 1)
        self.set_head_table(f"{datetime.datetime.now().strftime('%Y')}-W{datetime.datetime.now().strftime('%W')}")
        self.__oa = []
        self.__i = []
        self.__tst = []

    def to_json(self) -> str:
        json_dict = {
            "yot": self.__yot,
            "wotb": self.__wotb.isoformat(),
            "wote": self.__wote.isoformat(),
            "oa": [obj.__dict__ for obj in self.__oa],
            "i": [obj.__dict__ for obj in self.__i],
            "tst": [obj.__dict__ for obj in self.__tst]
        }
        return json.dumps(json_dict)

    @staticmethod
    def from_json(json_str: str) -> 'training_report':
        json_data = json.loads(json_str)
        report = training_report()
        report.__yot = json_data["yot"]
        report.__wotb = datetime.datetime.strptime(json_data["wotb"], "%Y-%m-%dT%H:%M:%S")
        report.__wote = datetime.datetime.strptime(json_data["wote"], "%Y-%m-%dT%H:%M:%S")

        report.__oa = []
        for item in json_data["oa"]:
            obj = tuple1(item["_tuple1__text"], item["_tuple1__hours"])
            report.__oa.append(obj)

        report.__i = []
        for item in json_data["i"]:
            obj = tuple1(item["_tuple1__text"], item["_tuple1__hours"])
            report.__i.append(obj)

        report.__tst = []
        for item in json_data["tst"]:
            obj = tuple1(item["_tuple1__text"], item["_tuple1__hours"])
            report.__tst.append(obj)

        return report

    def __calc_abj(self, wotb: datetime) -> int:
        diff = wotb.year - self.__training_begin.year
        if wotb.isocalendar()[1] > self.__training_begin.isocalendar()[1]:
            diff -= 1
        return diff

    def add_oa(self, text: str, hours: int):
        if len(self.__oa) <= 8:
            self.__oa.append(tuple1(text, hours))

    def edit_oa(self, index: int, hours: int):
        self.__oa[index].set_hours(hours)

    def remove_oa(self, index: int):
        self.__oa.pop(index)

    def add_i(self, text: str, hours: int):
        if len(self.__oa) <= 8:
            self.__i.append(tuple1(text, hours))

    def edit_i(self, index: int, hours: int):
        self.__i[index].set_hours(hours)

    def remove_i(self, index: int):
        self.__i.pop(index)

    def add_tst(self, text: str, hours: int):
        if len(self.__oa) <= 8:
            self.__tst.append(tuple1(text, hours))

    def edit_tst(self, index: int, hours: int):
        self.__tst[index].set_hours(hours)

    def remove_tst(self, index: int):
        self.__tst.pop(index)

    def __replace_markers(self, text: str) -> str:
        if text == "$YOT":
            return str(self.__yot)
        elif text == "$WOTB":
            return self.__wotb.strftime("%d.%m.%Y")
        elif text == "$WOTE":
            return self.__wote.strftime("%d.%m.%Y")
        elif text == "$OA":
            return tuple1.generate_str_from_list_text(self.__oa)
        elif text == "$OAH":
            return tuple1.generate_str_from_list_hours(self.__oa)
        elif text == "$I":
            return tuple1.generate_str_from_list_text(self.__i)
        elif text == "$IH":
            return tuple1.generate_str_from_list_hours(self.__i)
        elif text == "$TST":
            return tuple1.generate_str_from_list_text(self.__tst)
        elif text == "$TSTH":
            return tuple1.generate_str_from_list_hours(self.__tst)

    def set_head_table(self, cw: str):
        # kw must be format like jjjj-Www Example: "2023-W26"
        self.__wotb = datetime.datetime.strptime(cw + '-1', "%Y-W%W-%w")
        self.__wote = datetime.datetime.strptime(cw + '-5', "%Y-W%W-%w")
        self.__yot = self.__calc_abj(self.__wotb)

    def set_standard_oa(self):
        self.add_oa("RMM", int(input("Time for RMM: ")))
        self.add_oa("Helpdesk", int(input("Time for Helpdesk: ")))
        self.add_oa("Außendienst", int(input("Time for Außendienst: ")))
        self.add_oa("Softwareentwicklung", int(input("Time for Softwareentwicklung: ")))

    def set_standard_tst(self):
        self.add_tst("LF05: Programmieren", 4)
        self.add_tst("LF04: IT-Security", 6)
        self.add_tst("LF02: Arbeitsplätze ausstatten", 2)
        self.add_tst("LF01: Unternehmen und eigene Rolle", 2)
        self.add_tst("Englisch: Letter of complaint", 2)
        self.add_tst("Politikwissenschaft: Gesetze im Betrieblichen-Umfeld", 2)
        self.add_tst("Religion: Ethnische Diskussionen", 2)
        self.add_tst("Sport: Fallen gelernt", 4)

    def check_work_hours(self) -> (bool, int):
        work_hours = 0

        for obj in self.__oa:
            work_hours = work_hours + obj.get_hours()
        for obj in self.__i:
            work_hours = work_hours + obj.get_hours()
        for obj in self.__tst:
            work_hours = work_hours + obj.get_hours()

        if work_hours == 40:
            return True, work_hours
        else:
            return False, work_hours

    def save_document_to(self, filename: str):
        if self.check_work_hours()[0]:
            for table in self.__doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.startswith("$"):
                            cell.text = self.__replace_markers(cell.text)
            self.__doc.save(filename)
            self.__doc = Document("BerichtVorlage.docx")
        else:
            self.print_check_work_hours()

    def print_document(self):
        # Save the document as a temporary file
        temp_file = 'temp.docx'
        self.save_document_to(temp_file)

        # Print the document
        os.startfile(temp_file, "print")

    """
    const string TOP_LEFT_JOINT = "┌";
    const string TOP_RIGHT_JOINT = "┐";
    const string BOTTOM_LEFT_JOINT = "└";
    const string BOTTOM_RIGHT_JOINT = "┘";
    const string TOP_JOINT = "┬";
    const string BOTTOM_JOINT = "┴";
    const string LEFT_JOINT = "├";
    const string JOINT = "┼";
    const string RIGHT_JOINT = "┤";
    const char HORIZONTAL_LINE = '─';
    const char PADDING = ' ';
    const string VERTICAL_LINE = "│";
    """

    def print_head(self):
        print("┌" + "─Week of training Beginning──" +
              "┬" + "─Week of training Ending──" +
              "┬" + "─Year of training──" + "┐")
        print("│" + " " * 10 + self.__wotb.strftime("%d.%m.%Y") + " " * 9 +
              "│" + " " * 8 + self.__wote.strftime("%d.%m.%Y") + " " * 8 +
              "│" + " " * 9 + str(self.__yot) + " " * 9 + "│")
        print("└" + "─" * 29 + "┴" + "─" * 26 + "┴" + "─" * 19 + "┘")
        print()

    def print_tables(self):
        print("┌" + "─" + "Operational activities" + "─" * 42 + "┬──┐")
        for tup in self.__oa:
            print("│" + tup.get_text() + (" " * (65 - len(tup.get_text()))) + "│" +
                  (" " * (2 - len(str(tup.get_hours())))) + str(tup.get_hours()) + "│")
        print("└" + "─" * 65 + "┴──┘")
        print("┌" + "─" + "Instructions" + "─" * 52 + "┬──┐")
        for tup in self.__i:
            print("│" + tup.get_text() + (" " * (65 - len(tup.get_text()))) + "│" +
                  (" " * (2 - len(str(tup.get_hours())))) + str(tup.get_hours()) + "│")
        print("└" + "─" * 65 + "┴──┘")
        print("┌" + "─" + "Topics of school teaching" + "─" * 39 + "┬──┐")
        for tup in self.__tst:
            print("│" + tup.get_text() + (" " * (65 - len(tup.get_text()))) + "│" +
                  (" " * (2 - len(str(tup.get_hours())))) + str(tup.get_hours()) + "│")
        print("└" + "─" * 65 + "┴──┘")
        print()

    def print_check_work_hours(self):
        is_ok, work_hours = self.check_work_hours()
        status = 'ok' if is_ok else 'not ok'
        status_color = 'green' if is_ok else 'red'
        status_text = colored(status, status_color)
        print(f"\tWork hours are {status_text}:\t{work_hours}\n")

    def print_all(self):
        self.print_head()
        self.print_tables()
        self.print_check_work_hours()


def increment_cw(cw: str) -> str:
    split = cw.split("-W")
    split[1] = str(int(split[1]) + 1)
    return split[0] + "-W" + split[1]


def auto_generate_operational_week(cw: str) -> training_report:
    tr = training_report()
    tr.set_head_table(cw)
    tr.add_oa("RMM", random.randint(4, 12))
    tr.add_oa("Helpdesk", random.randint(4, 8))
    tr.add_oa("Außendienst", random.randint(2, 8))
    tr.add_oa("Softwareentwicklung", random.randint(8, 24))
    if tr.check_work_hours()[1] < 40:
        tr.add_oa("Betriebliche Tätigkeiten", random.randint(1, 5))
    return tr


def auto_generate_school_week(cw: str) -> training_report:
    tr = training_report()
    tr.set_head_table(cw)
    tr.add_oa("RMM", random.randint(2, 6))
    tr.add_oa("Helpdesk", random.randint(1, 3))
    tr.add_oa("Softwareentwicklung", random.randint(4, 10))
    if tr.check_work_hours()[1] < 16:
        tr.add_oa("Außendienst", random.randint(2, 5))
    if tr.check_work_hours()[1] < 16:
        tr.add_oa("Betriebliche Tätigkeiten", random.randint(1, 2))
    tr.set_standard_tst()
    return tr


def load_tr(path: str = "training_report.json") -> training_report:
    with open(path, "r") as file:
        json_str = file.read()
    return training_report.from_json(json_str)


def save_tr(training_rep: training_report, path: str = "training_report.json"):
    json_str = training_rep.to_json()
    with open(path, "w") as file:
        file.write(json_str)
    exit()


select_oa = """1 : set standard Operational activities
2 : add entry to Operational activities
3 : edit entry from Operational activities
4 : remove entry from Operational activities
0 : Back
"""


def operation_oa(tr: training_report):
    running = True
    while running:
        print(select_oa)
        select_action = input("Select action: ")
        print()
        if select_action == "0":
            # 0 : Back
            running = False
        elif select_action == "1":
            # 1 : set standard Operational activities
            tr.set_standard_oa()
            print()
        elif select_action == "2":
            # 2 : add entry to Operational activities
            tr.add_oa(input("Text: "), int(input("Hours: ")))
            print()
        elif select_action == "3":
            # 3 : edit entry from Operational activities
            tr.edit_oa(int(input("Line: ")) - 1, int(input("Hours: ")))
            print()
        elif select_action == "4":
            # 4 : remove entry from Operational activities
            tr.remove_oa(int(input("Line: ")) - 1)
            print()
        else:
            print(colored("Please Enter a valid action\n", "yellow"))


select_i = """1 : add entry to Instructions
2 : edit entry from Instructions
3 : remove entry from Instructions
0 : Back
"""


def operation_i(tr: training_report):
    running = True
    while running:
        print(select_i)
        select_action = input("Select action: ")
        print()
        if select_action == "0":
            # 0 : Back
            running = False
        elif select_action == "1":
            # 1 : add entry to Instructions
            tr.add_i(input("Text: "), int(input("Hours: ")))
            print()
        elif select_action == "2":
            # 2 : edit entry from Instructions
            tr.edit_i(int(input("Line: ")) - 1, int(input("Hours: ")))
            print()
        elif select_action == "3":
            # 3 : remove entry from Instructions
            tr.remove_i(int(input("Line: ")) - 1)
            print()
        else:
            print(colored("Please Enter a valid action\n", "yellow"))


select_tst = """1 : set standard Topics of school teaching
2 : add entry to Topics of school teaching
3 : edit entry from Topics of school teaching
4 : remove entry from Topics of school teaching
0 : Back
"""


def operation_tst(tr: training_report):
    running = True
    while running:
        print(select_tst)
        select_action = input("Select action: ")
        print()
        if select_action == "0":
            # 0 : Back
            running = False
        elif select_action == "1":
            # 1 : set standard Topics of school teaching
            tr.set_standard_tst()
            print()
        elif select_action == "2":
            # 2 : add entry to Topics of school teaching
            tr.add_tst(input("Text: "), int(input("Hours: ")))
            print()
        elif select_action == "3":
            # 3 : edit entry from Topics of school teaching
            tr.edit_tst(int(input("Line: ")) - 1, int(input("Hours: ")))
            print()
        elif select_action == "4":
            # 4 : remove entry from Topics of school teaching
            tr.remove_tst(int(input("Line: ")) - 1)
            print()
        else:
            print(colored("Please Enter a valid action\n", "yellow"))


def operation_th(tr: training_report):
    print("cw must be format like jjjj-Www Example: \"2023-W26\"")
    tr.set_head_table(input("cw: "))
    print()


select_io = """1 : save json
2 : save document
3 : load json
4 : load last training report
0 : Back
"""


def operation_io(tr: training_report) -> training_report:
    while True:
        print(select_io)
        select_action = input("Select action: ")
        print()
        if select_action == "0":
            # 0 : Back
            return tr
        elif select_action == "1":
            # 1 : save json
            save_tr(tr, input("Filename or Path: "))
            return tr
        elif select_action == "2":
            # 2 : save document
            tr.save_document_to(input("Filename or Path: "))
            return tr
        elif select_action == "3":
            # 3 : load json
            return load_tr(input("Filename or Path: "))
        elif select_action == "4":
            # 4 : load last training report
            return load_tr()
        else:
            print(colored("Please Enter a valid action\n", "yellow"))


def operation_print(tr: training_report):
    tr.print_document()


select_auto = """1 : generate operational week
2 : generate school week
0 : Back
"""


def operation_auto(tr: training_report) -> training_report:
    while True:
        print(select_auto)
        select_action = input("Select action: ")
        print()
        if select_action == "0":
            # 0 : Back
            return tr
        elif select_action == "1":
            # 1 : generate this week
            print("cw must be format like jjjj-Www Example: \"2023-W26\"")
            cw = input("cw: ")
            print()
            return auto_generate_operational_week(cw)
        elif select_action == "2":
            print("cw must be format like jjjj-Www Example: \"2023-W26\"")
            cw = input("cw: ")
            print()
            return auto_generate_school_week(cw)
        else:
            print(colored("Please Enter a valid action\n", "yellow"))


select_main = """1 : Operational activities
2 : Instructions
3 : Topics of school teaching
4 : Table Head
5 : I/O Operations
6 : Print Operations
7 : Automation
0 : Exit
"""


def main():
    tr = training_report()
    running = True

    while running:
        tr.print_all()
        print(select_main)
        select_action = input("Select action: ")
        print()
        if select_action == "0":
            # 0 : Exit
            running = False
        elif select_action == "1":
            # 1 : Operational activities
            operation_oa(tr)
        elif select_action == "2":
            # 2 : Instructions
            operation_i(tr)
        elif select_action == "3":
            # 3 : Topics of school teaching
            operation_tst(tr)
        elif select_action == "4":
            # 4 : Table Head
            operation_th(tr)
        elif select_action == "5":
            # 5 : I/O Operations
            tr = operation_io(tr)
        elif select_action == "6":
            # 6 : Print Operations
            operation_print(tr)
        elif select_action == "7":
            # 7 : Automation
            tr = operation_auto(tr)
        else:
            print(colored("Please Enter a valid action\n", "yellow"))

    save_tr(tr)


if __name__ == '__main__':
    print("\nWelcome to TrainingReportGenerator by Jannis Swientek\n")
    main()
